from docx.document import Document as DocxDocument
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from bs4.element import NavigableString
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX

def add_html_to_docx(doc: DocxDocument, html: str):

    highlights_map = {
        'yellow': 7,
        'red': 2,
        'green': 4,
        'cyan': 3,
        'blue': 5,
    }

    import re

    def rgb_to_hex(rgb_str):
        match = re.match(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', rgb_str)
        if match:
            r, g, b = match.groups()
            return '{:02X}{:02X}{:02X}'.format(int(r), int(g), int(b))
        return None

    def parse_style(style_str):
        styles = {}
        for part in style_str.split(';'):
            if ':' in part:
                key, val = part.split(':', 1)
                key = key.strip()
                val = val.strip()
                if key == 'color':
                    if val.startswith('rgb'):
                        hex_val = rgb_to_hex(val)
                        styles['color'] = hex_val if hex_val else val
                    else:
                        styles['color'] = val
                elif key == 'background-color':
                    if val.startswith('rgb'):
                        hex_val = rgb_to_hex(val)
                        styles['highlight'] = hex_val if hex_val else val
                    else:
                        styles['highlight'] = val
        return styles

    def set_run_background(run, rgb_color):
        rPr = run._r.get_or_add_rPr()
        # Remove existing shading elements if any
        for child in rPr.findall(qn('w:shd')):
            rPr.remove(child)
        shading_elm = parse_xml(
            r'<w:shd {} w:val="clear" w:color="auto" w:fill="{}"/>'.format(nsdecls('w'), rgb_color)
        )
        rPr.append(shading_elm)

    def apply_highlight_fallback(run, rgb_color):
        color_map = {
            'FFFF00': WD_COLOR_INDEX.YELLOW,
            'FF0000': WD_COLOR_INDEX.RED,
            '00FF00': WD_COLOR_INDEX.BRIGHT_GREEN,
            '00FFFF': WD_COLOR_INDEX.TURQUOISE,
            '0000FF': WD_COLOR_INDEX.BLUE,
            'FFC0CB': WD_COLOR_INDEX.PINK,
            '800080': WD_COLOR_INDEX.VIOLET,
            # add more if needed
        }
        rgb_color = rgb_color.upper()
        highlight_color = color_map.get(rgb_color, None)
        if highlight_color:
            run.font.highlight_color = highlight_color

    def apply_style_to_run(run, style):
        if style.get('bold'):
            run.bold = True
        if style.get('italic'):
            run.italic = True
        if style.get('underline'):
            run.underline = True
        # Strikethrough ignored as per user request
        if style.get('code'):
            run.font.name = 'Courier New'
        if 'color' in style:
            color_str = style['color']
            if color_str.startswith('#'):
                color_str = color_str[1:]
            try:
                run.font.color.rgb = RGBColor.from_string(color_str)
            except Exception:
                pass
        if 'highlight' in style:
            bg_color = style['highlight'].lstrip('#').upper()
            try:
                set_run_background(run, bg_color)
            except Exception:
                try:
                    apply_highlight_fallback(run, bg_color)
                except Exception:
                    pass

    def merge_styles(parent, new):
        """Merge two style dicts, with 'new' overriding."""
        merged = parent.copy()
        merged.update(new)
        return merged

    def extract_styles(el):
        """Extract styles from tag name and inline styles."""
        style = {}
        tag = el.name
        if tag in ['b', 'strong']:
            style['bold'] = True
        if tag in ['i', 'em']:
            style['italic'] = True
        if tag == 'u':
            style['underline'] = True
        if tag in ['s', 'strike', 'del']:
            style['strike'] = True
        if tag == 'code':
            style['code'] = True
        # Inline style parsing
        if 'style' in el.attrs:
            inline_styles = parse_style(el.attrs['style'])
            if 'color' in inline_styles:
                style['color'] = inline_styles['color']
            if 'highlight' in inline_styles:
                style['highlight'] = inline_styles['highlight']
        return style

    def process_element(el, paragraph, current_style):
        if isinstance(el, NavigableString):
            if el.strip():  
                run = paragraph.add_run(str(el))
                apply_style_to_run(run, current_style)
        else:
            el_style = extract_styles(el)
            combined_style = merge_styles(current_style, el_style)

            # Handle <a> specially to include link text and href
            if el.name == 'a':
                href = el.attrs.get('href', None)
                text_content = ''.join(str(c) for c in el.children if isinstance(c, NavigableString))
                run = paragraph.add_run()
                run.text = text_content
                apply_style_to_run(run, combined_style)
                return

            # Recursively process children
            for child in el.children:
                process_element(child, paragraph, combined_style)

    soup = BeautifulSoup(html, 'html.parser')

    for elem in soup.children:
        align = None
        classes = getattr(elem, 'attrs', {}).get('class', [])
        if 'ql-align-center' in classes:
            align = WD_ALIGN_PARAGRAPH.CENTER
        elif 'ql-align-right' in classes:
            align = WD_ALIGN_PARAGRAPH.RIGHT
        elif 'ql-align-left' in classes:
            align = WD_ALIGN_PARAGRAPH.LEFT

        if getattr(elem, "name", None) == 'p':
            p = doc.add_paragraph()
            if align is not None:
                p.alignment = align
            process_element(elem, p, {})
        elif getattr(elem, "name", None) in ['h1', 'h2', 'h3']:
            name = getattr(elem, "name", None)
            level = {'h1': 0, 'h2': 1, 'h3': 2}.get(name or '', 0)
            p = doc.add_heading(level=level)
            if align is not None:
                p.alignment = align
            process_element(elem, p, {})
        elif getattr(elem, "name", None) in ['ul', 'ol']:
            for li in getattr(elem, 'find_all', lambda *a, **k: [])('li'):
                p = doc.add_paragraph(style='List Bullet' if getattr(elem, "name", None) == 'ul' else 'List Number')
                if align is not None:
                    p.alignment = align
                process_element(li, p, {})
        elif isinstance(elem, NavigableString):
            if elem.strip():
                doc.add_paragraph(str(elem))